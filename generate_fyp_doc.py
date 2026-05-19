"""
FYP Documentation Generator — Orchestration Security Center
Generates FYP_Documentation.docx that consolidates existing .md docs AND
highlights NEW / UPDATED items (discovered from the actual codebase) in RED.

Color legend inside the docx:
    BLACK  -> content that already existed in the project's .md docs
    RED    -> NEW content / changes discovered by walking the live codebase
"""
from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Inches, Cm

# ─────────────────────────────────────────────────────────────────────────────
# CONFIG
# ─────────────────────────────────────────────────────────────────────────────

ROOT = Path(__file__).resolve().parent
OUTPUT = ROOT / "FYP_Documentation.docx"

RED   = RGBColor(0xC0, 0x00, 0x00)
BLACK = RGBColor(0x00, 0x00, 0x00)
GRAY  = RGBColor(0x55, 0x55, 0x55)
BLUE  = RGBColor(0x1F, 0x49, 0x7D)

BODY_FONT = "Times New Roman"
CODE_FONT = "Consolas"

# Track everything we mark in red for the changelog summary at the end
CHANGELOG: list[tuple[str, str]] = []  # (section, item)


def log_change(section: str, item: str) -> None:
    CHANGELOG.append((section, item))


# ─────────────────────────────────────────────────────────────────────────────
# DOCX HELPERS
# ─────────────────────────────────────────────────────────────────────────────

def add_run(paragraph, text, *, bold=False, italic=False, color=BLACK,
            font=BODY_FONT, size=12):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = font
    run.font.size = Pt(size)
    run.font.color.rgb = color
    return run


def add_paragraph(doc, text="", *, red=False, bold=False, italic=False,
                  size=12, justify=True, indent=None):
    p = doc.add_paragraph()
    if justify:
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        p.paragraph_format.left_indent = indent
    p.paragraph_format.space_after = Pt(4)
    if text:
        add_run(p, text, bold=bold, italic=italic,
                color=(RED if red else BLACK), size=size)
    return p


def add_mixed(doc, segments, *, justify=True):
    """segments: list of (text, red_bool) tuples for inline color mixing."""
    p = doc.add_paragraph()
    if justify:
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(4)
    for text, red in segments:
        add_run(p, text, color=(RED if red else BLACK))
    return p


def add_heading(doc, text, *, level=1, red=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    sizes = {0: 26, 1: 18, 2: 14, 3: 12}
    size = sizes.get(level, 12)
    run = add_run(p, text, bold=True, color=(RED if red else BLUE), size=size)
    # Use a heading style so a TOC could be regenerated in Word later
    try:
        p.style = doc.styles[f"Heading {level}"]
        # Re-apply our color because the style resets it
        for r in p.runs:
            r.font.color.rgb = (RED if red else BLUE)
            r.font.name = BODY_FONT
            r.font.size = Pt(size)
            r.bold = True
    except KeyError:
        pass
    return p


def add_bullet(doc, segments, *, indent_level=0):
    """segments: list of (text, red_bool)."""
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 + 0.25 * indent_level)
    p.paragraph_format.space_after = Pt(2)
    for text, red in segments:
        add_run(p, text, color=(RED if red else BLACK))
    return p


def add_code_block(doc, code, *, red=False):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._element.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F4F4F4")
    pPr.append(shd)
    add_run(p, code, color=(RED if red else BLACK), font=CODE_FONT, size=9)
    return p


def add_page_break(doc):
    doc.add_page_break()


def add_table(doc, headers, rows, *, red_row_mask=None, red_cell_mask=None):
    """
    rows: list of list[str].
    red_row_mask: list[bool] same length as rows — whole row in red.
    red_cell_mask: list[list[bool]] — per-cell coloring (overrides row mask).
    """
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    # Header
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        add_run(p, h, bold=True, color=BLACK, size=11)
        hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # Body
    for ri, row in enumerate(rows):
        row_red = bool(red_row_mask and ri < len(red_row_mask) and red_row_mask[ri])
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = ""
            cell_red = row_red
            if red_cell_mask and ri < len(red_cell_mask) and ci < len(red_cell_mask[ri]):
                cell_red = red_cell_mask[ri][ci]
            add_run(cells[ci].paragraphs[0], str(val),
                    color=(RED if cell_red else BLACK), size=10)
    return table


def set_default_style(doc):
    style = doc.styles["Normal"]
    style.font.name = BODY_FONT
    style.font.size = Pt(12)
    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)


def add_page_numbers(doc):
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run.font.name = BODY_FONT
    run.font.size = Pt(10)


# ─────────────────────────────────────────────────────────────────────────────
# SECTION BUILDERS
# ─────────────────────────────────────────────────────────────────────────────

def build_title_page(doc):
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "FINAL YEAR PROJECT REPORT", bold=True, size=22, color=BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "\nOrchestration Security Center", bold=True, size=28)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p,
            "An AI-Assisted Deterministic Cybersecurity Orchestration "
            "Platform for Small and Medium Enterprises",
            italic=True, size=14, color=GRAY)

    for _ in range(6):
        doc.add_paragraph()

    info_rows = [
        ("Team Leader", "Omar Kapil"),
        ("Team Size", "11 members across 4 sub-teams"),
        ("Project Type", "AI-Driven DAST + SOC Orchestration Platform"),
        ("Submission Date", "July 2, 2026"),
        ("Document Generated", datetime.now().strftime("%Y-%m-%d %H:%M")),
        ("Document Version", "2.0  (red = new since last .md docs)"),
    ]
    t = doc.add_table(rows=len(info_rows), cols=2)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for (k, v), row in zip(info_rows, t.rows):
        row.cells[0].text = ""
        row.cells[1].text = ""
        add_run(row.cells[0].paragraphs[0], k, bold=True, size=12)
        add_run(row.cells[1].paragraphs[0], v, size=12)

    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p,
            "Color legend in this document:  "
            "BLACK = content from existing project docs  ·  ",
            italic=True, size=10, color=GRAY)
    add_run(p, "RED = new or updated since the last .md documentation",
            italic=True, size=10, color=RED)
    add_page_break(doc)


def build_abstract(doc):
    add_heading(doc, "Abstract", level=1)
    add_paragraph(doc,
        "Orchestration Security Center is a deterministic cybersecurity "
        "orchestration platform that gives Small and Medium Enterprises (SMEs) "
        "professional-grade threat visibility without requiring a dedicated "
        "Security Operations Centre. The system chains industry-standard "
        "open-source tools — Nmap, Nuclei, OpenVAS, Wazuh, Elasticsearch and "
        "n8n — through a four-stage rule-based pipeline (Recon → Targeted "
        "Chaining → Validation → Risk Scoring), translates the technical "
        "findings into a unified 0–100 Risk Score and 0–100 Health Score, "
        "and surfaces them in a real-time React dashboard backed by FastAPI, "
        "Celery, Redis and PostgreSQL.")
    add_paragraph(doc,
        "AI (Google Gemini 1.5 Flash) is used strictly as a Technical "
        "Educator: it explains findings and writes remediation guidance, "
        "but never makes scoring or scanning decisions, preserving "
        "reproducibility and auditability.")
    add_mixed(doc, [
        ("This second-iteration release expands the platform from a "
         "single-user prototype into a hardened multi-user system: ", False),
        ("a full JWT authentication stack with three RBAC roles "
         "(VIEWER/ANALYST/ADMIN), tamper-evident hash-chained audit logs, "
         "cryptographically signed PDF reports, CVSS-driven risk breakdowns, "
         "finding deduplication, SLA clocks, framework control tagging "
         "(OWASP/CWE/ISO/NIST/PCI-DSS), scope allowlists with per-target "
         "rate limits, scan checkpointing and orphan reaping, runtime "
         "feature-flag overrides, and an expanded ten-container, "
         "four-subnet lab environment.", True),
    ])
    log_change("Abstract", "Iteration-2 hardening summary (auth, RBAC, audit chain, "
                           "signed reports, CVSS, SLA, framework tags, scope guard).")
    add_page_break(doc)


def build_toc(doc):
    add_heading(doc, "Table of Contents", level=1)
    toc = [
        ("1.  Introduction & Problem Statement", False),
        ("2.  Objectives & Scope", False),
        ("3.  Background & Related Work", False),
        ("4.  System Architecture", False),
        ("5.  Technology Stack & Dependencies", False),
        ("6.  Backend — Modules, Services & Endpoints", False),
        ("7.  Database Schema", False),
        ("8.  Frontend — Pages, Components & State", False),
        ("9.  Authentication, RBAC & Audit Logging", True),
        ("10. Lab Environment", False),
        ("11. The 4-Stage Scan Pipeline", False),
        ("12. Risk Scoring, CVSS, SLA & Framework Tagging", True),
        ("13. AI Advisory Role", False),
        ("14. SIEM & SOAR Integration", False),
        ("15. Reliability, Hardening & Audit Chain", True),
        ("16. Use Cases & Target Audience", False),
        ("17. Advantages & Disadvantages", False),
        ("18. Evidence, Testing & Evaluation", False),
        ("19. Team Structure & Roles", False),
        ("20. Setup, Installation & Deployment", False),
        ("21. Future Work", False),
        ("22. Conclusion", False),
        ("Appendix A. Changelog of Updates (Red Items)", True),
    ]
    for title, red in toc:
        add_paragraph(doc, title, red=red, justify=False)
    add_page_break(doc)


def build_intro(doc):
    add_heading(doc, "1.  Introduction & Problem Statement", level=1)
    add_paragraph(doc,
        "Small and Medium Enterprises face a structural protection gap: "
        "they cannot justify the budget of a full Security Operations "
        "Centre, yet they rely on the same internet-facing infrastructure "
        "as large enterprises. Off-the-shelf scanners produce hundreds of "
        "raw alerts with little business context; CVSS scores do not "
        "translate into the language of priority that a non-specialist IT "
        "administrator can act on; and stitching together Nmap, Nuclei, "
        "OpenVAS, a SIEM and a SOAR engine is a full-time job.")
    add_paragraph(doc,
        "Orchestration Security Center is the answer to that gap. The "
        "name is a play on the HTTP 404 status — the system finds the "
        "vulnerabilities other tools miss or bury in noise. Instead of "
        "firing every test at every target, it uses a rule-based chaining "
        "engine: when Nmap finds port 445 open, it runs SMB-specific "
        "Nuclei templates, not web SQL injection tests. The end result is "
        "that one thousand raw logs are reduced to five prioritised, "
        "plain-language action items.")


def build_objectives(doc):
    add_heading(doc, "2.  Objectives & Scope", level=1)
    add_paragraph(doc,
        "The project pursues six concrete objectives:")
    add_bullet(doc, [("Unify network scanning, vulnerability detection, SIEM, "
                      "SOAR and AI advisory under one dashboard.", False)])
    add_bullet(doc, [("Translate raw findings into business-readable risk and "
                      "health scores.", False)])
    add_bullet(doc, [("Keep every scoring and scanning decision deterministic "
                      "for reproducibility.", False)])
    add_bullet(doc, [("Use AI strictly in an advisory role — explanation and "
                      "remediation, never decision-making.", False)])
    add_bullet(doc, [("Deploy the entire stack with a single docker compose "
                      "command.", False)])
    add_bullet(doc, [("Harden the platform for multi-user operation with "
                      "JWT authentication, three-role RBAC, hash-chained "
                      "audit logging and cryptographically signed reports.",
                      True)])
    log_change("Objectives", "Multi-user hardening objective (auth, RBAC, "
                             "audit chain, signed reports).")


def build_background(doc):
    add_heading(doc, "3.  Background & Related Work", level=1)
    add_paragraph(doc,
        "Commercial competitors such as Tenable Nessus (USD 3,990 per "
        "year per scanner) and Burp Suite Enterprise (USD 8,999 per year) "
        "offer comparable scanning depth but at a price point that "
        "excludes most SMEs. Open-source alternatives — Nmap, Nuclei, "
        "OpenVAS, Wazuh, n8n — are individually powerful but require "
        "significant integration effort and cybersecurity expertise to "
        "operate as a single workflow. Orchestration Security Center "
        "fills that integration gap with a deterministic chaining engine, "
        "a unified data model and a non-specialist dashboard.")


def build_architecture(doc):
    add_heading(doc, "4.  System Architecture", level=1)
    add_paragraph(doc,
        "The platform is a multi-tier containerised system. Each tier is "
        "an independent service that communicates over well-defined "
        "protocols — REST and WebSocket between browser and API, Redis "
        "pub/sub between API and Celery workers, and SQLAlchemy ORM "
        "between application code and PostgreSQL.")
    add_code_block(doc, """\
┌─────────────────────────────────────────────────────────────┐
│  USER BROWSER  —  React 18 + Vite  (localhost:5173)         │
└──────────────────┬──────────────────────────────────────────┘
                   │ HTTPS + WSS (via Caddy / Nginx TLS proxy)
┌──────────────────▼──────────────────────────────────────────┐
│  FASTAPI BACKEND  —  Uvicorn (localhost:8000)               │
│   /api/v1/*    ── REST endpoints                             │
│   /ws/logs     ── WebSocket fan-out                          │
│   /health      ── liveness / readiness                       │
│   AgentOrchestrator  ── 4-stage deterministic pipeline       │
└──────────────────┬──────────────────────────────────────────┘
                   │ async tasks via Celery
┌──────────────────▼──────────────────────────────────────────┐
│  TASK QUEUE LAYER  —  Redis broker + Celery workers          │
└──────────────────┬──────────────────────────────────────────┘
                   │
┌──────────────────▼──────────────────────────────────────────┐
│  SCANNING TOOLS  —  Nmap, Nuclei, OpenVAS GVM, validation   │
│                     probe, UnifiedRiskEngine                 │
└──────────────────┬──────────────────────────────────────────┘
                   │
┌──────────────────▼──────────────────────────────────────────┐
│  DATA LAYER  —  PostgreSQL 15 (prod)  ·  SQLite (dev)        │
└──────────────────┬──────────────────────────────────────────┘
                   │ events forwarded to
┌──────────────────▼──────────────────────────────────────────┐
│  SIEM  —  Wazuh  →  Elasticsearch  →  Kibana                 │
│  SOAR —  n8n (workflow automation, remediation playbooks)    │
└─────────────────────────────────────────────────────────────┘
""")
    add_paragraph(doc,
        "The reverse-proxy layer (Caddy or Nginx) terminates TLS and "
        "routes /api and /ws to the FastAPI service and / to the Vite "
        "build, allowing the entire deployment to live behind a single "
        "HTTPS endpoint.",
        red=True)
    log_change("Architecture",
               "TLS reverse-proxy layer (Caddy/Nginx) terminating HTTPS/WSS "
               "in front of FastAPI.")


def build_tech_stack(doc):
    add_heading(doc, "5.  Technology Stack & Dependencies", level=1)
    add_paragraph(doc,
        "The platform is built entirely on open-source components. Items "
        "in red were introduced in the current iteration and are not "
        "covered by the previous Markdown stack table.")

    add_heading(doc, "5.1  Backend dependencies (Python 3.10)", level=2)
    rows = [
        ("fastapi", "REST + WebSocket framework", False),
        ("uvicorn[standard]", "ASGI server", False),
        ("sqlalchemy + alembic", "ORM + migrations", False),
        ("aioredis / redis", "Async Redis client / Celery broker", False),
        ("celery", "Async task queue", False),
        ("httpx", "Async HTTP client", False),
        ("psycopg2-binary / asyncpg / aiosqlite", "Postgres + SQLite drivers", False),
        ("python-nmap", "Nmap subprocess wrapper", False),
        ("python-gvm", "OpenVAS / GVM client", False),
        ("google-genai >= 0.8.0", "Gemini AI advisory client", False),
        ("reportlab", "PDF report generation", False),
        ("pydantic-settings + python-dotenv", "Typed settings & env loading", False),
        ("python-jose[cryptography] >= 3.3.0", "JWT signing/verification (Phase 3.1)", True),
        ("passlib[bcrypt] + bcrypt==4.0.1", "Password hashing (Phase 3.1)", True),
        ("cryptography >= 42.0.0", "Credential & report-signing primitives (Phase 3.2)", True),
        ("croniter == 3.0.3", "Scheduled scan cron parsing", True),
        ("python-multipart >= 0.0.9", "Avatar upload via POST /auth/me/avatar", True),
        ("aiolimiter >= 1.1.0", "Per-target outbound rate limiting", True),
    ]
    add_table(doc, ["Package", "Purpose"],
              [(r[0], r[1]) for r in rows],
              red_row_mask=[r[2] for r in rows])
    for r in rows:
        if r[2]:
            log_change("Backend dependencies", f"{r[0]} — {r[1]}")

    add_heading(doc, "5.2  Frontend dependencies (Node 20 / Vite 5)", level=2)
    rows = [
        ("react @ 18.2", "UI runtime", False),
        ("vite @ 5", "Dev server + build", False),
        ("tailwindcss", "Utility CSS framework", False),
        ("axios", "REST client", False),
        ("recharts", "Charting", False),
        ("@tanstack/react-query @ 5", "Server-state caching & invalidation", True),
        ("@tanstack/react-query-devtools", "Query inspector in dev", True),
        ("zustand", "Lightweight global store", True),
        ("react-router-dom @ 7", "Client-side routing (Login, Settings, etc.)", True),
        ("framer-motion", "Animations / page transitions", True),
        ("ldrs", "Loader / spinner primitives", True),
        ("react-window", "List virtualisation", True),
        ("react-force-graph-2d", "Force-directed topology graph", True),
        ("d3-force / d3-hierarchy / d3-selection / d3-transition / d3-ease",
         "Modular D3 sub-packages replacing the d3 monolith", True),
        ("lucide-react", "Icon set", True),
        ("vitest + @testing-library/* + msw + jsdom", "Unit/component test stack", True),
    ]
    add_table(doc, ["Package", "Purpose"],
              [(r[0], r[1]) for r in rows],
              red_row_mask=[r[2] for r in rows])
    for r in rows:
        if r[2]:
            log_change("Frontend dependencies", f"{r[0]} — {r[1]}")


def build_backend(doc):
    add_heading(doc, "6.  Backend — Modules, Services & Endpoints", level=1)

    add_heading(doc, "6.1  Directory layout", level=2)
    add_code_block(doc, """\
backend/
├── app/
│   ├── main.py                # FastAPI app + lifespan + WS endpoint
│   ├── core/
│   │   ├── config.py          # Pydantic Settings + runtime overrides
│   │   ├── database.py        # SQLAlchemy engine + SessionLocal
│   │   ├── security.py        # bcrypt + JWT helpers          (NEW)
│   │   ├── audit.py           # log_action() helper           (NEW)
│   │   ├── request_id.py      # RequestIdMiddleware           (NEW)
│   │   └── celery_app.py      # Celery configuration
│   ├── api/
│   │   ├── api.py             # router aggregator
│   │   ├── deps.py            # get_current_user, require_role
│   │   └── v1/endpoints/
│   │       ├── auth.py             (NEW — JWT login/logout/me)
│   │       ├── rbac.py             (NEW — admin user mgmt)
│   │       ├── audit.py            (NEW — hash-chained log)
│   │       ├── findings.py         (NEW — dedup'd findings)
│   │       ├── config.py           (NEW — runtime feature flags)
│   │       ├── dashboard.py
│   │       ├── scans.py
│   │       ├── targets.py
│   │       ├── vulnerabilities.py
│   │       ├── network.py
│   │       ├── openvas.py
│   │       ├── reports.py
│   │       ├── siem.py
│   │       └── lab.py
│   ├── models/
│   │   ├── scan.py            # Target, Scan, Vulnerability, Finding,
│   │   │                      #   AgentLog, Endpoint, Report, ScanAsset,
│   │   │                      #   AssetService, NetworkAsset, ActionItem
│   │   ├── user.py            (NEW — User + UserRole enum)
│   │   ├── audit_log.py       (NEW — AuditLog)
│   │   └── config.py          (NEW — RuntimeConfig)
│   ├── schemas/               # Pydantic request / response models
│   └── services/              # see §6.3
├── alembic/versions/          # 4 -> 14 migration files
├── scripts/
│   ├── full_system_check.py
│   └── simulate_attack.py
├── tests/                     # pytest suite (e2e, risk, integration)
├── requirements.txt
├── Dockerfile
└── alembic.ini
""", red=False)

    add_heading(doc, "6.2  REST endpoints", level=2)
    add_paragraph(doc,
        "Base URL (dev): http://localhost:8000/api/v1   ·   "
        "Base URL (prod via reverse proxy): https://localhost/api/v1   ·   "
        "Interactive Swagger: /docs", italic=True)
    add_paragraph(doc,
        "All routes except /health and /auth/login require an "
        "Authorization: Bearer <JWT> header. Tokens expire after 30 "
        "minutes and are obtained via POST /auth/login.")

    endpoint_rows = [
        ("Auth", "POST /auth/login", "Issue JWT — returns { access_token, role, force_password_change }", True),
        ("Auth", "POST /auth/logout", "Revoke session (204)", True),
        ("Auth", "GET  /auth/me", "Current user profile", True),
        ("Auth", "PATCH /auth/me", "Self-edit full_name / bio / phone", True),
        ("Auth", "POST /auth/me/avatar", "Multipart avatar upload", True),
        ("Auth", "POST /auth/change-password", "{ old_password, new_password }", True),
        ("RBAC (admin)", "GET    /rbac/users", "List users (paginated)", True),
        ("RBAC (admin)", "POST   /rbac/users", "Create user with role", True),
        ("RBAC (admin)", "PATCH  /rbac/users/{id}/role", "Promote / demote", True),
        ("RBAC (admin)", "POST   /rbac/users/{id}/disable", "Soft-disable login", True),
        ("RBAC (admin)", "POST   /rbac/users/{id}/enable", "Re-enable login", True),
        ("RBAC (admin)", "POST   /rbac/users/{id}/reset-password", "Admin force-reset", True),
        ("RBAC (admin)", "DELETE /rbac/users/{id}", "Hard delete", True),
        ("RBAC (admin)", "GET    /rbac/audit-logs", "Recent audit entries", True),
        ("Targets", "GET    /targets/", "List scope-managed targets", False),
        ("Targets", "POST   /targets/", "Create target (admin)", False),
        ("Targets", "GET    /targets/{id}", "Detail", False),
        ("Targets", "DELETE /targets/{id}", "Remove (admin)", False),
        ("Scans", "GET  /scans/", "Paginated list", False),
        ("Scans", "POST /scans/", "Create scan { target_id, scan_type, tools?, schedule? }", False),
        ("Scans", "GET  /scans/{id}", "Full detail incl. findings", False),
        ("Scans", "GET  /scans/{id}/task-status", "Celery state", False),
        ("Scans", "POST /scans/{id}/stop", "Cancel in-flight scan", False),
        ("Scans", "GET  /scans/{id}/audit/verify", "Verify hash-chained AgentLog", True),
        ("Findings", "GET  /findings/", "Deduplicated findings (filter: severity, scan_id)", True),
        ("Findings", "GET  /findings/{id}", "Full finding + remediation + control_tags", True),
        ("Findings", "POST /findings/{id}/false-positive", "Mark as FP", True),
        ("Vulnerabilities", "GET /vulnerabilities/", "Raw vuln observations", False),
        ("Reports", "GET  /reports/", "List reports", False),
        ("Reports", "POST /reports/", "{ scan_id, format } — pdf | json | html", False),
        ("Reports", "GET  /reports/{id}/download", "Signed PDF download", False),
        ("Reports", "GET  /reports/{id}/verify", "Verify HMAC signature", True),
        ("Network", "GET /network/assets", "Discovered assets", False),
        ("Network", "GET /network/assets/{id}", "Asset detail + timeline", False),
        ("Network", "GET /network/topology", "D3-ready force-graph payload", False),
        ("SIEM", "GET /siem/alerts", "Recent Wazuh alerts (paginated)", False),
        ("SIEM", "GET /siem/correlations", "Alerts ↔ findings linkage", False),
        ("Audit", "GET /audit/log", "Paginated hash-chained audit", True),
        ("Audit", "GET /audit/verify", "Confirm chain integrity", True),
        ("Lab", "GET  /lab/services", "Health of lab containers", False),
        ("Lab", "POST /lab/scenarios/{name}/run", "Run a scripted scenario", False),
        ("OpenVAS (admin)", "POST /openvas/tasks", "Deep-scan task", False),
        ("OpenVAS (admin)", "GET  /openvas/tasks/{id}", "Status + findings", False),
        ("Config (admin)", "GET  /config/", "Runtime feature flags", True),
        ("Config (admin)", "PATCH /config/{key}", "Toggle / update flag at runtime", True),
        ("System", "GET /health", "Liveness + readiness + schema drift", False),
        ("System", "WSS /ws/logs", "Live event fan-out", False),
    ]
    add_table(doc,
        ["Group", "Method + Path", "Purpose"],
        [(r[0], r[1], r[2]) for r in endpoint_rows],
        red_row_mask=[r[3] for r in endpoint_rows])
    for r in endpoint_rows:
        if r[3]:
            log_change("Backend endpoints", f"{r[1]} — {r[2]}")

    add_heading(doc, "6.3  Service layer (backend/app/services/)", level=2)
    services = [
        ("agent_orchestrator.py", "Core 4-stage pipeline engine — Recon → Targeted Chaining → Validation → Risk Scoring. Uses SERVICE_TO_TEMPLATE map to deterministically route findings.", False),
        ("unified_risk_engine.py", "Converts vulnerabilities into Risk Score (0–100) and Health Score (0–100) using CVSS severity × asset criticality.", False),
        ("scan_tasks.py", "Celery task definitions; wraps the orchestrator in asyncio.new_event_loop() to avoid loop conflicts inside workers.", False),
        ("discovery_agent.py", "Nmap-based network reconnaissance.", False),
        ("nmap_wrapper.py", "Subprocess wrapper around nmap CLI.", False),
        ("nuclei_wrapper.py", "Subprocess wrapper around nuclei CLI.", False),
        ("openvas.py", "OpenVAS / GVM GMP client.", False),
        ("intelligence_agent.py", "Gemini-based AI advisory (advisory only).", False),
        ("ai_advisor.py", "Pluggable LLM advisor abstraction sitting above intelligence_agent; lets the IntelligenceAgent be swapped for a local model without touching callers.", True),
        ("elastic_integration.py", "Elasticsearch query layer.", False),
        ("wazuh_integration.py", "Wazuh REST API client.", False),
        ("infrastructure_agent.py", "Infrastructure health monitor.", False),
        ("asset_monitor.py", "Asset tracking & alerting.", False),
        ("soar_orchestrator.py", "Triggers n8n remediation workflows.", False),
        ("pdf_generator.py", "Builds the security assessment PDF via ReportLab.", False),
        ("lab_manager.py", "Manages lab container lifecycle.", False),
        ("ws_manager.py", "WebSocket connection registry + broadcast helpers.", False),
        ("event_publisher.py", "Publishes worker events on the Redis ws_events channel.", False),
        ("validation_probe.py", "Phase 1.3 — independent reprobe that confirms a Nuclei finding before it is allowed to influence the risk score. LLM justification (validation_notes) is logged but never overrides the probe verdict.", True),
        ("cvss.py", "Phase 4.1 — CVSS v3.1 vector parser + base-score calculator; populates Vulnerability.cvss_score and Vulnerability.cvss_vector.", True),
        ("finding_dedup.py", "Phase 4.2 — computes the SHA-256 fingerprint (target_id + vuln_type + normalised_url + parameter + evidence-match) and either creates a new Finding or links the observation to an existing one.", True),
        ("framework_tagger.py", "Phase 5.3 — maps Nuclei template categories to OWASP Top-10, CWE, ISO 27001 Annex A, NIST CSF and PCI-DSS control IDs; stored in Finding.control_tags.", True),
        ("llm_guard.py", "Safety wrapper around all Gemini calls — strips PII, rate-limits, validates JSON shape, and falls back to deterministic text on any guardrail trip.", True),
        ("report_signer.py", "Phase 5.2 — HMAC-SHA256 signs PDF bytes with REPORT_SIGNING_KEY and stores the signature in the Report row for tamper detection.", True),
        ("sla.py", "Phase 4.3 — SLA clock: assigns Finding.due_date from severity-based policy (e.g. CRITICAL = 7 days) and exposes breach metrics.", True),
        ("alert_correlator.py", "Cross-references incoming Wazuh alerts against open Findings to expose /siem/correlations.", True),
        ("scoring_explainer.py", "Produces the human-readable Scan.risk_breakdown JSON consumed by the RiskBreakdownDrawer component.", True),
        ("task_monitor.py", "Celery worker health probe used by /health and the Settings panel.", True),
        ("scope_guard.py", "Phase 2.4 — enforces Target.scope_allowlist (hostnames + CIDRs) and per-target max_rps / max_concurrent_scans before any outbound packet leaves the worker.", True),
        ("scan_reaper.py", "Phase 2 — runs on lifespan startup to mark any scan stuck in RUNNING from before the restart as FAILED with failure_reason='orphaned_on_restart'.", True),
        ("topology_generator.py", "Builds the D3 force-graph payload (nodes + links + zone metadata) served by /network/topology.", True),
    ]
    for name, desc, red in services:
        add_mixed(doc, [(name + " — ", red), (desc, red)])
        if red:
            log_change("Backend services", f"{name} — {desc.split('.')[0]}.")


def build_db_schema(doc):
    add_heading(doc, "7.  Database Schema", level=1)
    add_paragraph(doc,
        "PostgreSQL 15 in production, SQLite in development. SQLAlchemy "
        "ORM with Alembic migrations. The schema is split across the core "
        "scan domain (scan.py), the authentication domain (user.py), the "
        "audit domain (audit_log.py) and the feature-flag domain "
        "(config.py).")

    add_heading(doc, "7.1  Tables", level=2)
    tables = [
        ("targets", "Scope-managed scan target (URL or CIDR).", False),
        ("scans", "One scan session per row.", False),
        ("vulnerabilities", "Per-observation findings of a scan.", False),
        ("scan_assets / asset_services", "Discovered hosts and their services.", False),
        ("network_assets", "Persistent inventory across scans.", False),
        ("action_items", "Action items generated by the risk engine.", False),
        ("agent_logs", "AI agent reasoning trail, now hash-chained.", False),
        ("endpoints", "Discovered API endpoints (URL + method + params).", False),
        ("findings", "Phase 4.2 — deduplicated, persistent issue per target (one Finding ↔ many Vulnerability observations).", True),
        ("reports", "Phase 5.2 — generated PDF + findings_hash + HMAC signature + stored pdf_bytes for re-verification.", True),
        ("users", "Phase 3.1 — authentication identities with role, profile (full_name/bio/phone/avatar_url), disabled flag, force_password_change.", True),
        ("audit_logs", "Phase 5.1 — append-only audit of every privileged action (actor_email, action, target_id, detail).", True),
        ("runtime_config", "Phase B — key/value feature flag overrides applied over env defaults at startup and on PATCH /config/{key}.", True),
    ]
    add_table(doc, ["Table", "Purpose"],
              [(t[0], t[1]) for t in tables],
              red_row_mask=[t[2] for t in tables])
    for t in tables:
        if t[2]:
            log_change("DB tables", f"{t[0]} — {t[1].split('—')[0].strip()}")

    add_heading(doc, "7.2  Notable column additions", level=2)
    add_paragraph(doc, "Target table additions:")
    add_bullet(doc, [("scope_allowlist (JSON) — list of hostnames / CIDRs the scanner may touch.", True)])
    add_bullet(doc, [("max_rps (int, default 10) — per-target outbound request ceiling.", True)])
    add_bullet(doc, [("max_concurrent_scans (int, default 1) — Redis-locked concurrency cap; second scan returns failure_reason='concurrency_limit'.", True)])
    add_bullet(doc, [("environment_type (lab|development|staging|production) — gates aggressiveness and auto-report.", True)])
    add_bullet(doc, [("compliance_tags (JSON) — e.g. [\"pci-dss\",\"hipaa\",\"iso-27001\",\"gdpr\"].", True)])
    add_bullet(doc, [("notes (text), last_scanned_at (datetime) — operator metadata.", True)])

    add_paragraph(doc, "Scan table additions:")
    add_bullet(doc, [("risk_breakdown (JSON) — structured CVSS breakdown produced by UnifiedRiskEngine.calculate_scan_risk_v2().", True)])
    add_bullet(doc, [("failure_reason (str) — orphaned_on_restart | concurrency_limit | <exception message>.", True)])
    add_bullet(doc, [("checkpoint (str) — recon_done | attack_done | validated | risk_scored | reported; skipped by run_full_scan on Celery retry.", True)])
    add_bullet(doc, [("environment_type — inherited from target at scan creation.", True)])

    add_paragraph(doc, "Vulnerability table additions:")
    add_bullet(doc, [("cvss_vector + cvss_score (Phase 4.1).", True)])
    add_bullet(doc, [("finding_id FK (Phase 4.2) — links observation to the dedup'd Finding.", True)])
    add_bullet(doc, [("validation_notes (Phase 1.3) — LLM justification, never overrides reprobe.", True)])
    add_bullet(doc, [("raw_request, raw_response, evidence_hash, detected_by, template_id (Phase 1.2 Nuclei evidence).", True)])

    add_paragraph(doc, "AgentLog tamper-evidence (Phase 5.1):")
    add_bullet(doc, [("prev_hash + this_hash — SHA-256 hash chain; verified by GET /scans/{id}/audit/verify.", True)])

    for k in ["scope_allowlist", "max_rps", "max_concurrent_scans",
              "environment_type", "compliance_tags", "risk_breakdown",
              "failure_reason", "checkpoint", "cvss_score", "finding_id",
              "validation_notes", "Nuclei evidence fields", "hash-chained AgentLog"]:
        log_change("DB columns", k)

    add_heading(doc, "7.3  RBAC roles", level=2)
    add_table(doc,
        ["Role", "Permissions"],
        [
            ("VIEWER", "Read-only — GET every resource, no mutations."),
            ("ANALYST", "VIEWER + create / stop scans + update vulnerability status + mark false-positives."),
            ("ADMIN", "ANALYST + target CRUD + user / role management + runtime config toggles."),
        ],
        red_row_mask=[True, True, True])
    log_change("RBAC", "VIEWER / ANALYST / ADMIN role hierarchy.")


def build_frontend(doc):
    add_heading(doc, "8.  Frontend — Pages, Components & State", level=1)

    add_heading(doc, "8.1  Pages", level=2)
    pages = [
        ("Dashboard.jsx", "Primary tab-controller page (Overview, Vulnerabilities, Topology, SIEM, Reports, Lab).", False),
        ("LoginPage.jsx", "JWT login form with force-password-change redirect.", True),
        ("SignUpPage.jsx", "Self-service registration (admin-approval workflow).", True),
        ("SettingsPage.jsx", "Runtime feature-flag toggles + profile / SLA / scope panels.", True),
        ("UserManagementPage.jsx", "Admin-only user list with role change, disable / enable, reset-password, delete.", True),
        ("ProfilePage.jsx", "Self-edit full_name / bio / phone + avatar upload via POST /auth/me/avatar.", True),
    ]
    for n, d, r in pages:
        add_mixed(doc, [(n + " — ", r), (d, r)])
        if r:
            log_change("Frontend pages", f"{n} — {d.split('.')[0]}.")

    add_heading(doc, "8.2  Component groups", level=2)
    add_paragraph(doc,
        "frontend/src/components/ is organised into four buckets: top-level "
        "shared components, dashboard/ panels rendered by Dashboard.jsx, "
        "ui/ primitives, and OpenVAS/ scanner-specific widgets.")

    add_paragraph(doc, "Dashboard panels (frontend/src/components/dashboard/):", bold=True)
    panels = [
        ("StatCards", "Top-row KPI metric cards.", False),
        ("ScanButton", "Scan initiation + pipeline progress.", False),
        ("ScanConfigModal", "Per-scan tool / aggressiveness picker before launch.", True),
        ("QuickScanPopover", "Inline 'scan this asset' shortcut from any table row.", True),
        ("NetworkTopology", "D3 force-graph of discovered assets.", False),
        ("TopologyLegend", "Severity / zone legend for the topology view.", True),
        ("ExposureMap", "Zone-by-zone exposure heatmap.", True),
        ("VulnerabilitiesPanel", "Filterable vulnerability table.", False),
        ("VulnTrend", "Historical severity trend chart.", False),
        ("SeverityDonut", "Severity-distribution donut chart.", True),
        ("RiskHeatmap", "Asset × severity matrix.", False),
        ("RiskScore", "Gauge ring for current scan risk score.", False),
        ("RiskBreakdownDrawer", "CVSS-aware drill-down on Scan.risk_breakdown.", True),
        ("RemediationPanel", "Per-finding remediation guidance, SLA clock & owner.", True),
        ("UptimeGauge", "Health-score gauge.", False),
        ("ActionCenter", "Action items list.", False),
        ("ActivityFeed", "Recent events (scans, role changes, alerts).", False),
        ("AgentLogViewer", "Per-stage agent reasoning chain.", False),
        ("OrchestrationFeed", "Real-time pipeline event stream.", False),
        ("ScanPipelinePanel", "4-stage pipeline visualiser.", False),
        ("ScanHistory", "Past scans table.", False),
        ("ScanningBanner", "Sticky banner while a scan is RUNNING.", True),
        ("TargetsManager", "Target CRUD + scope-allowlist editor.", False),
        ("EnvironmentWizard", "First-run wizard that captures environment_type, compliance_tags and scope allowlist.", True),
        ("LabEnvironment", "Lab container status.", False),
        ("Reports", "Report list + verify-signature button.", False),
        ("AssetDetailPanel", "Per-asset detail + AI advice.", False),
        ("AssetTimeline", "Time-series of detections per asset.", True),
        ("IncidentDetailDrawer", "Drill-down on SIEM incidents.", False),
        ("UnifiedInbox", "Notifications inbox aggregating WS events.", False),
        ("LiveConsole", "Real-time worker stdout.", False),
        ("SettingsPanel", "In-dashboard runtime feature-flag editor (mirrors SettingsPage).", True),
        ("Taskbar", "Bottom dock showing in-flight scans and last 3 toasts.", True),
    ]
    for n, d, r in panels:
        add_mixed(doc, [("· " + n + " — ", r), (d, r)])
        if r:
            log_change("Frontend dashboard panels", f"{n} — {d.split('.')[0]}.")

    add_paragraph(doc, "UI primitives (frontend/src/components/ui/):", bold=True)
    prims = [
        ("CyberButton, CyberBadge", "Themed buttons / badges.", False),
        ("GaugeRing", "Reusable gauge primitive.", False),
        ("Tabs, SubTabBar", "Tab containers.", False),
        ("SkeletonPulse", "Loading skeleton.", False),
        ("Toast", "Toast primitive used by ToastProvider.", False),
        ("ProtectedRoute", "react-router guard — redirects unauthenticated users to /login.", True),
        ("RoleGuard", "Wraps children in a role check (e.g. <RoleGuard role='ADMIN'>).", True),
        ("RoleBadge", "Coloured badge for VIEWER / ANALYST / ADMIN.", True),
        ("ConfirmDialog", "Reusable confirmation modal.", True),
        ("EmptyState", "Empty-state illustration + CTA.", True),
    ]
    for n, d, r in prims:
        add_mixed(doc, [("· " + n + " — ", r), (d, r)])
        if r:
            log_change("Frontend UI primitives", f"{n} — {d.split('.')[0]}.")

    add_paragraph(doc, "Top-level components:", bold=True)
    tops = [
        ("ReportGenerator, SecurityAdvisor, MetricCard, DeviceDetailModal, ErrorBoundary, ToastProvider, TabNavigation, VulnerabilityList, Dashboard",
         "Pre-existing top-level components.", False),
        ("CommandPalette", "Ctrl-K command palette for power-user navigation.", True),
        ("ShortcutCheatsheet", "Modal listing all keyboard shortcuts.", True),
        ("NotificationsBell", "Top-bar bell with unread-count badge driven by WS events.", True),
    ]
    for n, d, r in tops:
        add_mixed(doc, [("· " + n + " — ", r), (d, r)])
        if r:
            log_change("Frontend top-level components", f"{n} — {d.split('.')[0]}.")

    add_heading(doc, "8.3  State management", level=2)
    add_paragraph(doc,
        "Pre-existing: AuthContext.jsx (authentication state) and "
        "RealTimeContext.jsx (WebSocket reducer; handles KPI_UPDATE, "
        "SCAN_STATUS, LOG_MESSAGE, ORCHESTRATION_EVENT, ASSET_UPDATE).")
    add_paragraph(doc,
        "New: zustand stores for ephemeral UI state, "
        "@tanstack/react-query for server-state caching and automatic "
        "invalidation after mutations, react-router-dom v7 for "
        "client-side routing across the new Login / Settings / Profile / "
        "UserManagement pages.", red=True)
    log_change("Frontend state", "zustand + @tanstack/react-query + react-router-dom v7 stack.")


def build_auth_rbac(doc):
    add_heading(doc, "9.  Authentication, RBAC & Audit Logging", level=1, red=True)
    add_paragraph(doc,
        "The platform previously ran in single-user mode with no login. "
        "This iteration introduces a complete authentication and "
        "authorisation stack — every claim about single-user operation in "
        "older documentation is now obsolete.", red=True)

    add_heading(doc, "9.1  Authentication flow", level=2, red=True)
    add_code_block(doc, """\
User → POST /auth/login { email, password }
   ↓
FastAPI verifies bcrypt password_hash, signs a JWT (HS256, 30 min TTL)
   ↓
Frontend stores token in memory + axios interceptor attaches
       Authorization: Bearer <jwt> to every subsequent request
   ↓
ProtectedRoute (react-router) redirects to /login if token missing/expired
   ↓
RoleGuard hides / disables UI affordances the user lacks the role for
   ↓
Backend dependency require_role(UserRole.ANALYST | ADMIN) returns 403
       when an authenticated request lacks the required role
""", red=True)

    add_heading(doc, "9.2  Roles", level=2, red=True)
    add_paragraph(doc, "VIEWER — read-only across all resources.", red=True)
    add_paragraph(doc, "ANALYST — VIEWER plus scan triggering, vulnerability "
                       "status updates, false-positive marking.", red=True)
    add_paragraph(doc, "ADMIN — ANALYST plus target CRUD, full user / role "
                       "management, runtime feature-flag toggles, "
                       "audit-log access.", red=True)

    add_heading(doc, "9.3  Default admin seed", level=2, red=True)
    add_paragraph(doc,
        "On first boot the application checks for a user with email "
        "admin@local and, if absent, seeds one with password Admin#159 "
        "and force_password_change=True. The administrator is forced to "
        "rotate the password before any other action is allowed.", red=True)
    log_change("Auth", "Default admin seed (admin@local / Admin#159, force change).")

    add_heading(doc, "9.4  Profile & avatar", level=2, red=True)
    add_paragraph(doc,
        "Users have self-service profile fields (full_name, bio, phone, "
        "avatar_url) edited via PATCH /auth/me. Avatars are uploaded as "
        "multipart/form-data to POST /auth/me/avatar and served from "
        "/api/v1/avatars (a StaticFiles mount under settings.API_V1_STR), "
        "so the same reverse-proxy rule that routes the API also handles "
        "avatar requests.", red=True)

    add_heading(doc, "9.5  Hash-chained audit log", level=2, red=True)
    add_paragraph(doc,
        "Every privileged action (user create/disable/delete, role "
        "change, password reset, scan launch, target mutation, runtime "
        "config change, report generation) is appended to the audit_logs "
        "table by app.core.audit.log_action() with { actor_id, "
        "actor_email, action, target_id, detail, created_at }.", red=True)
    add_paragraph(doc,
        "AI agent reasoning is independently chained inside agent_logs: "
        "each row carries prev_hash and this_hash = SHA-256(prev_hash + "
        "canonical_json(scan_id, agent_name, action, reasoning)). "
        "GET /scans/{id}/audit/verify replays the chain and returns "
        "{ ok: bool, broken_at: id|null } — making the agent trace "
        "tamper-evident.", red=True)
    log_change("Audit", "Hash-chained AgentLog with /scans/{id}/audit/verify.")
    log_change("Audit", "audit_logs table + log_action() for every privileged action.")


def build_lab(doc):
    add_heading(doc, "10.  Lab Environment", level=1)
    add_paragraph(doc,
        "The lab is a self-contained, multi-subnet network of "
        "intentionally vulnerable Docker containers that simulates a "
        "realistic SME enterprise topology. It now spans four subnets "
        "and ten active containers across DMZ, Corporate, Data and "
        "Management zones.")

    add_heading(doc, "10.1  Lab target personas", level=2)
    targets = [
        ("dmz", "lab_webserver", "webserver.sme-lab.local", "OWASP Juice Shop (:3000)", "SQLi, XSS, BOLA, IDOR, broken-auth, SSRF", "9.5"),
        ("dmz", "lab_api_gateway", "api-gw.sme-lab.local", "Nginx legacy API (:8081)", "Info disclosure, header leak, directory listing, exposed Swagger", "6.0"),
        ("dmz", "lab_dns_server", "dns.sme-lab.local", "CoreDNS (:5353)", "Zone transfer, DNS amplification", "5.0"),
        ("corp", "lab_fileserver", "fileserver.sme-lab.local", "Samba (:1139,:4445)", "Weak credentials, SMB enum, HR data exposure", "8.0"),
        ("corp", "lab_mailserver", "mail.sme-lab.local", "GreenMail (:3025,:3110,:3143,:8082)", "Plaintext protocols, weak creds, user enum", "7.0"),
        ("corp", "lab_workstation", "ws01.sme-lab.local", "Nginx HR workstation (:8083)", "Info disclosure, internal-network leak", "4.0"),
        ("data", "lab_database", "db.sme-lab.local", "PostgreSQL 13 (:5433)", "Weak password, no SSL, sensitive PII", "9.0"),
        ("data", "lab_redis_cache", "cache.sme-lab.local", "Redis 6 (:6380)", "No auth, protected-mode off, cross-subnet reachable", "8.5"),
    ]
    add_table(doc,
        ["Zone", "Container", "Hostname", "Service / Ports", "Vulnerabilities", "CVSS"],
        targets)

    add_heading(doc, "10.2  Lab support services", level=2)
    add_table(doc, ["Zone", "Container", "Purpose"], [
        ("mgmt", "lab_traffic_gen", "Generates realistic background traffic across all lab subnets for SIEM data."),
        ("mgmt", "lab_log_shipper", "Ships lab events and traffic logs to Elasticsearch and Wazuh."),
    ])

    add_heading(doc, "10.3  Network topology", level=2)
    add_code_block(doc, """\
the-dashboard-project_lab_network (external bridge)
│
├── dmz subnet  (10.10.10.0/24)
│   ├── lab_webserver    10.10.10.10  → :3000   (Juice Shop)
│   ├── lab_api_gateway  10.10.10.20  → :8081   (Nginx legacy API)
│   └── lab_dns_server   10.10.10.30  → :5353   (CoreDNS)
│
├── corp subnet (10.10.20.0/24)
│   ├── lab_fileserver   10.10.20.10  → :1139, :4445  (Samba)
│   ├── lab_mailserver   10.10.20.20  → :3025, :3110, :3143, :8082
│   └── lab_workstation  10.10.20.40  → :8083
│
├── data subnet (10.10.30.0/24)
│   ├── lab_database     10.10.30.10  → :5433   (PostgreSQL)
│   └── lab_redis_cache  10.10.30.20  → :6380   (Redis)
│
└── mgmt subnet (10.10.40.0/24)
    ├── lab_traffic_gen  10.10.40.10
    └── lab_log_shipper  10.10.40.20
""")
    add_paragraph(doc,
        "An isolation override (infra/isolation/docker-compose.lab."
        "isolation.override.yml) can be layered on top of the standard "
        "lab compose file to drop inter-subnet routing for negative "
        "tests during evidence runs.", red=True)
    log_change("Lab", "Isolation override compose file for inter-subnet "
                      "negative testing.")


def build_pipeline(doc):
    add_heading(doc, "11.  The 4-Stage Scan Pipeline", level=1)
    add_code_block(doc, """\
Stage 1 — RECON   (ReconAgent + Nmap)
   Target → Nmap scan → discovered ports + services
   → stored in ScanAsset / AssetService

Stage 2 — TARGETED CHAINING  (AttackAgent + Nuclei)
   For each discovered service:
       port 80/443   → Nuclei web templates (SQLi, XSS, BOLA)
       port 445      → Nuclei SMB (enum, default-login)
       port 22 / 21  → Nuclei SSH / FTP default credentials
       port 6379     → Nuclei Redis unauthenticated access
       port 8080+    → Nuclei HTTP API misconfiguration

Stage 3 — VALIDATION  (confidence ≥ 0.6 + reprobe)
   validation_probe.py independently confirms each Nuclei finding;
   LLM validation_notes are logged but never override the probe.

Stage 4 — RISK SCORING  (UnifiedRiskEngine)
   - Risk Score   (0–100) = Σ severity_weight × occurrence × asset_criticality
   - Health Score (0–100) = 100 − deductions
   - ActionItem generation: REMEDIATION | REVIEW | CONFIGURATION
   - Finding deduplication: each observation linked to a persistent
     Finding row by SHA-256 fingerprint
   - CVSS breakdown stored in Scan.risk_breakdown
   - SLA clock: Finding.due_date derived from severity
   - Framework tags: OWASP / CWE / ISO27001 / NIST / PCI-DSS
   - Results persisted + broadcast via WebSocket → live dashboard update
""")
    add_mixed(doc, [
        ("Reliability additions: ", True),
        ("each phase writes a Scan.checkpoint value (recon_done | "
         "attack_done | validated | risk_scored | reported), so a Celery "
         "retry skips phases already completed. On startup, scan_reaper "
         "marks any scan stuck in RUNNING as FAILED with "
         "failure_reason='orphaned_on_restart'. The scope_guard "
         "enforces Target.scope_allowlist and per-target max_rps / "
         "max_concurrent_scans before any outbound packet leaves the "
         "worker.", True),
    ])
    log_change("Pipeline", "Scan checkpointing + scan_reaper + scope_guard "
                           "reliability layer.")


def build_scoring(doc):
    add_heading(doc, "12.  Risk Scoring, CVSS, SLA & Framework Tagging",
                level=1, red=True)

    add_heading(doc, "12.1  Risk Score formula", level=2)
    add_code_block(doc, """\
Risk Score = Σ (severity_weight × occurrence) × asset_criticality_multiplier

Severity weights:    CRITICAL=10.0  HIGH=7.0  MEDIUM=4.0  LOW=1.5  INFO=0.1
Asset multipliers:   database_server=1.5x  web_server=1.3x  workstation=1.0x
Output is clamped to 0–100.
""")

    add_heading(doc, "12.2  Health Score formula", level=2)
    add_code_block(doc, """\
Health Score = 100
   − (critical_count × 25)
   − (high_count     × 15)
   − dangerous_port_penalties
Clamped to 0–100.
""")

    add_heading(doc, "12.3  Action item rules", level=2)
    add_table(doc, ["Condition", "Action Type", "Priority"], [
        ("Critical or High vulnerability", "REMEDIATION", "Immediate"),
        ("Medium vulnerability", "REVIEW", "Scheduled"),
        ("Dangerous open port (22, 445, 3389, 6379)", "CONFIGURATION", "High"),
    ])

    add_heading(doc, "12.4  CVSS-aware risk breakdown (Phase 4.1)",
                level=2, red=True)
    add_paragraph(doc,
        "Every Vulnerability now carries cvss_vector (e.g. CVSS:3.1/"
        "AV:N/AC:L/PR:N/UI:N/S:U/C:H/I:H/A:H) and cvss_score (parsed "
        "base 0–10). UnifiedRiskEngine.calculate_scan_risk_v2() writes a "
        "structured risk_breakdown JSON onto Scan.risk_breakdown:", red=True)
    add_code_block(doc, """\
{
  "score": 74.3,
  "breakdown": [
    { "vuln_id": "…", "title": "…", "cvss": 9.8,
      "severity": "critical", "asset_criticality": 1.5,
      "contribution": 24.5 },
    …
  ]
}
""", red=True)
    add_paragraph(doc,
        "The RiskBreakdownDrawer component renders this payload, "
        "letting the operator see exactly which findings drove the "
        "score.", red=True)

    add_heading(doc, "12.5  Finding deduplication (Phase 4.2)",
                level=2, red=True)
    add_paragraph(doc,
        "finding_dedup.py computes fingerprint = SHA-256(target_id + "
        "vuln_type + normalised_url + parameter + evidence_match_"
        "signature), enforced unique per (target_id, fingerprint). "
        "A Vulnerability observation either links to an existing Finding "
        "(re-detection updates last_seen) or creates a new one. Status "
        "lifecycle: OPEN → FIXED | ACCEPTED | REOPENED | FALSE_POSITIVE.",
        red=True)

    add_heading(doc, "12.6  SLA clock (Phase 4.3)", level=2, red=True)
    add_paragraph(doc,
        "sla.py assigns Finding.due_date from a severity-based policy "
        "(CRITICAL = 7 days, HIGH = 30, MEDIUM = 90 by default) when the "
        "Finding is created and exposes breach metrics on the dashboard. "
        "Owners are tracked in Finding.owner_user_id.", red=True)

    add_heading(doc, "12.7  Framework control tags (Phase 5.3)",
                level=2, red=True)
    add_paragraph(doc,
        "framework_tagger.py maps Nuclei template categories to "
        "Finding.control_tags:", red=True)
    add_code_block(doc, """\
{
  "owasp_top10":         "A03:2021",
  "cwe":                 "CWE-89",
  "iso27001_annex_a":    "A.12.6.1",
  "nist_csf_function":   "PR.IP",
  "pci_dss_requirement": "6.3.1"
}
""", red=True)
    add_paragraph(doc,
        "Empty dictionary when the template category is unknown — tags "
        "are never invented.", red=True)
    log_change("Scoring", "Phase 4.1 CVSS breakdown, Phase 4.2 dedup, "
                          "Phase 4.3 SLA, Phase 5.3 framework tags.")


def build_ai(doc):
    add_heading(doc, "13.  AI Advisory Role", level=1)
    add_paragraph(doc,
        "Gemini 1.5 Flash is used strictly as a Technical Educator. It "
        "does NOT make scan decisions. After the deterministic engine "
        "has scored each top-3 critical asset, IntelligenceAgent "
        "produces:")
    add_code_block(doc, """\
{
  "risk_explanation":   "Why this vulnerability is dangerous",
  "business_impact":    "What it means for your business",
  "remediation_advice": "Exact steps to fix it",
  "response_priority":  "immediate | scheduled | monitor"
}
""")
    add_paragraph(doc,
        "Output is rendered in the AssetDetailPanel under 'SME Security "
        "Advisor'. Without a Gemini API key the system falls back to "
        "generic advisory text — no functionality is lost.")
    add_paragraph(doc,
        "llm_guard.py wraps every Gemini call: it strips PII, enforces a "
        "global request budget via aiolimiter, validates the response "
        "matches the expected JSON shape, and substitutes a "
        "deterministic fallback whenever any guardrail trips.", red=True)
    log_change("AI", "llm_guard.py PII-stripping + rate-limit + JSON-shape "
                     "guard around Gemini calls.")


def build_siem(doc):
    add_heading(doc, "14.  SIEM & SOAR Integration", level=1)
    add_paragraph(doc, "Wazuh (EDR / IDS):", bold=True)
    add_bullet(doc, [("Receives agent events from lab containers via lab_log_shipper.", False)])
    add_bullet(doc, [("Detects suspicious activity patterns (brute force, port scans).", False)])
    add_bullet(doc, [("Alerts forwarded into the dashboard via wazuh_integration.py.", False)])

    add_paragraph(doc, "Elasticsearch + Kibana:", bold=True)
    add_bullet(doc, [("Receives forwarded logs from Wazuh and direct lab container output.", False)])
    add_bullet(doc, [("Powers the SIEM Events tab. Kibana available at localhost:5601.", False)])

    add_paragraph(doc, "n8n (SOAR):", bold=True)
    add_bullet(doc, [("Available at localhost:5678 with pre-configured remediation playbooks.", False)])
    add_bullet(doc, [("Triggered by ActionItem events from the backend.", False)])

    add_paragraph(doc, "Alert correlation:", bold=True)
    add_bullet(doc, [("alert_correlator.py cross-references incoming Wazuh alerts against open Findings; "
                      "/siem/correlations surfaces the linkage in the UI.", True)])
    log_change("SIEM", "alert_correlator.py + /siem/correlations endpoint.")


def build_hardening(doc):
    add_heading(doc, "15.  Reliability, Hardening & Audit Chain",
                level=1, red=True)
    add_paragraph(doc,
        "The current iteration introduced a dedicated reliability and "
        "hardening track. Each item below addresses a class of failures "
        "or attack that older versions could not detect or recover from.",
        red=True)

    add_paragraph(doc, "Phase 1.2/1.3 — Evidence + reprobe validation:", bold=True, red=True)
    add_paragraph(doc,
        "Nuclei evidence captured into raw_request, raw_response, "
        "evidence_hash, detected_by, template_id. validation_probe.py "
        "reproves every finding before it is allowed into the risk "
        "engine; LLM validation_notes never override the probe.",
        red=True)

    add_paragraph(doc, "Phase 2 — Scan reliability:", bold=True, red=True)
    add_paragraph(doc,
        "Scan.checkpoint lets a Celery retry skip already-completed "
        "phases. scan_reaper marks scans stuck in RUNNING from before a "
        "restart as FAILED (failure_reason='orphaned_on_restart'). "
        "Scan.failure_reason captures the first 120 chars of any "
        "unexpected exception for forensics.", red=True)

    add_paragraph(doc, "Phase 2.4 — Scope guard:", bold=True, red=True)
    add_paragraph(doc,
        "Target.scope_allowlist (hostnames + CIDRs), max_rps and "
        "max_concurrent_scans are enforced by scope_guard.py before any "
        "outbound packet leaves the worker. A second concurrent scan on "
        "the same target is rejected with "
        "failure_reason='concurrency_limit'.", red=True)

    add_paragraph(doc, "Phase 3.1 — JWT authentication:", bold=True, red=True)
    add_paragraph(doc,
        "python-jose + passlib[bcrypt]; tokens expire after 30 minutes; "
        "force_password_change flag on first login; default admin@local "
        "seeded with Admin#159.", red=True)

    add_paragraph(doc, "Phase 3.2 — Credential encryption:", bold=True, red=True)
    add_paragraph(doc,
        "Target.auth_credentials is encrypted at rest using the "
        "cryptography library so JWT and basic-auth secrets cannot be "
        "extracted from a stolen database dump.", red=True)

    add_paragraph(doc, "Phase 5.1 — Hash-chained AgentLog:", bold=True, red=True)
    add_paragraph(doc,
        "Each AgentLog row carries prev_hash and this_hash; "
        "GET /scans/{id}/audit/verify replays the chain and surfaces "
        "any tampering point.", red=True)

    add_paragraph(doc, "Phase 5.2 — Signed PDF reports:", bold=True, red=True)
    add_paragraph(doc,
        "report_signer.py HMAC-SHA256 signs PDF bytes with "
        "REPORT_SIGNING_KEY. The signature and the canonical "
        "findings_hash are stored on the Report row; "
        "GET /reports/{id}/verify confirms the bytes have not been "
        "altered.", red=True)

    add_paragraph(doc, "Operational additions:", bold=True, red=True)
    add_bullet(doc, [("RequestIdMiddleware adds an X-Request-ID header to every response for log correlation.", True)])
    add_bullet(doc, [("Auto-Alembic migrations on lifespan startup (alembic upgrade head).", True)])
    add_bullet(doc, [("Health endpoint reports schema-drift status (current vs head revision), cached 5 min to avoid Alembic log flooding.", True)])
    add_bullet(doc, [("Redis pub/sub listener with exponential backoff (2s → 32s) and quiet-period DEBUG logging.", True)])
    add_bullet(doc, [("Avatar StaticFiles mount under /api/v1/avatars so the existing reverse-proxy rule covers it.", True)])

    log_change("Hardening", "Phase 1.2/1.3 evidence + reprobe.")
    log_change("Hardening", "Phase 2 scan checkpoints + reaper.")
    log_change("Hardening", "Phase 2.4 scope guard + rate limits.")
    log_change("Hardening", "Phase 3.1 JWT + bcrypt auth.")
    log_change("Hardening", "Phase 3.2 credential encryption.")
    log_change("Hardening", "Phase 5.1 hash-chained AgentLog.")
    log_change("Hardening", "Phase 5.2 signed PDF reports + /reports/{id}/verify.")
    log_change("Hardening", "RequestIdMiddleware + X-Request-ID correlation.")
    log_change("Hardening", "Auto-Alembic migrations + schema-drift health check.")


def build_use_cases(doc):
    add_heading(doc, "16.  Use Cases & Target Audience", level=1)
    add_paragraph(doc,
        "Orchestration Security Center turns thousands of raw security "
        "alerts into roughly five prioritised, plain-language action "
        "items that tell a non-technical person exactly what to fix and "
        "how.")

    add_heading(doc, "16.1  Target audience", level=2)
    add_table(doc, ["Audience", "How they use it"], [
        ("SME IT Administrators", "One-click security assessments without cybersecurity expertise."),
        ("Small Business Owners", "Health-score dashboard for at-a-glance posture."),
        ("Non-specialist IT teams", "AI-generated step-by-step remediation guidance."),
        ("Cybersecurity students & educators", "Lab environment for safe practice."),
        ("Security consultants", "Rapid SME assessments with PDF deliverables."),
        ("DevOps engineers", "Docker Compose deployment for continuous monitoring."),
    ])

    add_heading(doc, "16.2  Real-world use cases", level=2)
    cases = [
        ("SME security posture assessment",
         "Small e-commerce company scans web server + internal network; gets a clear action item to parameterise SQL queries in the checkout endpoint."),
        ("Network infrastructure auditing",
         "Mid-sized IT department audits internal subnets; Nmap discovers hosts, Nuclei checks for weak credentials on SMB and Redis."),
        ("Compliance preparation",
         "Annual security review evidence: comprehensive signed PDF report of all findings, severity, and remediation steps taken."),
        ("Educational lab environment",
         "Students practise on the built-in 10-container, 4-subnet lab."),
        ("Continuous monitoring",
         "Startup integrates with Wazuh + n8n; critical vulnerabilities auto-trigger remediation playbooks."),
        ("Incident response support",
         "IT team scans an affected segment, identifies exposed services, prioritises fixes via AI-generated business impact analysis."),
    ]
    for title, desc in cases:
        add_paragraph(doc, f"• {title} — {desc}")


def build_pros_cons(doc):
    add_heading(doc, "17.  Advantages & Disadvantages", level=1)

    add_heading(doc, "17.1  Advantages", level=2)
    add_paragraph(doc, "1.  All-in-one platform — Nmap, Nuclei, OpenVAS, Wazuh + Elastic, n8n and Gemini under one dashboard.")
    add_paragraph(doc, "2.  Zero cybersecurity expertise required for the operator.")
    add_paragraph(doc, "3.  Deterministic and reproducible scoring decisions.")
    add_paragraph(doc, "4.  AI advisory without AI risk — Gemini explains but never decides.")
    add_paragraph(doc, "5.  100% open source; no licensing fees.")
    add_paragraph(doc, "6.  One-command deployment (docker compose up -d).")
    add_paragraph(doc, "7.  Real-time WebSocket dashboard with sub-second latency.")
    add_paragraph(doc, "8.  Smart vulnerability chaining (service → matched templates only).")
    add_paragraph(doc, "9.  Built-in lab environment for safe practice.")
    add_paragraph(doc, "10. PDF report generation for stakeholders.", )
    add_paragraph(doc, "11. Scalable microservices architecture.")
    add_paragraph(doc, "12. Comprehensive risk scoring (severity × occurrence × asset criticality × exposure × port penalties).")
    add_paragraph(doc, "13. Multi-user with three-role RBAC (VIEWER / ANALYST / ADMIN); admin-managed user lifecycle; force-rotate seeded admin password.", red=True)
    add_paragraph(doc, "14. Tamper-evident: hash-chained AgentLog + HMAC-signed PDF reports + audit_logs of every privileged action.", red=True)
    add_paragraph(doc, "15. CVSS-aware risk breakdown with per-finding contribution drill-down.", red=True)
    add_paragraph(doc, "16. SLA clock per Finding with severity-derived due_date and ownership tracking.", red=True)
    add_paragraph(doc, "17. Compliance-ready: each Finding tagged with OWASP Top-10, CWE, ISO 27001 Annex A, NIST CSF, PCI-DSS control IDs.", red=True)
    add_paragraph(doc, "18. Safety net: scope_allowlist + max_rps + max_concurrent_scans block out-of-scope or accidentally aggressive scans before a packet leaves the worker.", red=True)
    add_paragraph(doc, "19. Crash-safe: scan checkpoints + orphan reaper + per-scan failure_reason.", red=True)
    add_paragraph(doc, "20. Runtime feature flags via /config/{key} without backend restart.", red=True)
    log_change("Advantages", "Items 13–20 added (RBAC, audit chain, signed reports, "
                              "CVSS breakdown, SLA, framework tags, scope guard, "
                              "checkpoints, runtime config).")

    add_heading(doc, "17.2  Disadvantages", level=2)
    add_paragraph(doc, "1.  No authenticated scanning — cannot crawl behind login forms.")
    add_paragraph(doc, "2.  ", )
    p = doc.paragraphs[-1]
    add_run(p, "Originally listed as 'single-user / no RBAC' — ", italic=True, color=GRAY)
    add_run(p, "this limitation has been removed in the current iteration (see §9).", color=RED, italic=True)
    add_paragraph(doc, "3.  Validated only in lab / Docker environments; real SME networks (firewalls, NAT, VPN, dynamic IPs) not extensively tested.")
    add_paragraph(doc, "4.  High system resource requirements — ≥8 GB RAM for the full stack.")
    add_paragraph(doc, "5.  Limited concurrent scanning (Celery is sequential per-target by default).")
    add_paragraph(doc, "6.  AI advisory requires a Google Gemini API key for personalised output (deterministic fallback otherwise).")
    add_paragraph(doc, "7.  ", )
    p = doc.paragraphs[-1]
    add_run(p, "SIEM / SOAR integration partially implemented — ", italic=True, color=GRAY)
    add_run(p, "alert correlation has landed (alert_correlator.py + /siem/correlations); end-to-end Wazuh-agent auto-provisioning remains manual.", color=RED, italic=True)
    add_paragraph(doc, "8.  No mobile companion app — web only.")
    add_paragraph(doc, "9.  ", )
    p = doc.paragraphs[-1]
    add_run(p, "Originally listed as 'no compliance mapping' — ", italic=True, color=GRAY)
    add_run(p, "control tags (OWASP/CWE/ISO/NIST/PCI) are now stored per Finding; full per-framework report templates remain future work.", color=RED, italic=True)
    add_paragraph(doc, "10. No SAST — Orchestration Security Center remains a DAST platform.")
    add_paragraph(doc, "11. Internet dependency for the Gemini AI features (advisory only).")
    add_paragraph(doc, "12. Docker dependency — the whole stack is container-native.")
    log_change("Disadvantages",
               "Items 2, 7 and 9 superseded by RBAC, alert correlation, "
               "and framework tagging respectively.")


def build_evidence(doc):
    add_heading(doc, "18.  Evidence, Testing & Evaluation", level=1)
    add_paragraph(doc,
        "Evidence runs are organised under the evidence/ directory by "
        "phase. Each phase folder contains a run_summary.md plus any "
        "supporting heuristic / data-model / isolation / failure-mode / "
        "accuracy reports.")

    add_table(doc, ["Phase folder", "Contents"], [
        ("evidence/phase2/", "run_summary.md — initial reliability evidence."),
        ("evidence/phase3/", "heuristic_audit.md — heuristic scoring audit."),
        ("evidence/phase4/", "data_model_review.md + run_summary.md — CVSS / dedup / SLA evidence."),
        ("evidence/phase5/", "isolation_audit.md, failure_modes.md, run_summary.md — isolation override + failure mode evidence."),
        ("evidence/phase6/", "accuracy_report.md, run_summary.md — accuracy evaluation."),
        ("evidence/phases_4_5_6_sign_off.md", "Combined sign-off."),
        ("docs/audit/baseline_2026-04-24.md", "Baseline audit."),
    ], red_row_mask=[True, True, True, True, True, True, True])
    log_change("Evidence", "evidence/phase{2..6} + sign-off + audit baseline.")

    add_paragraph(doc, "Automated test suites:", bold=True)
    add_bullet(doc, [("backend/tests/ — pytest: test_e2e_scans.py, test_risk.py, test_risk_engine_manual.py.", False)])
    add_bullet(doc, [("Frontend: vitest + @testing-library + msw mocks (see package.json scripts: test / test:ui / test:coverage).", True)])

    add_paragraph(doc, "Lab attack scenarios (lab/scenarios/):", bold=True)
    add_bullet(doc, [("misconfig_scenario.md — service misconfiguration scenario.", False)])
    add_bullet(doc, [("sqli_scenario.md — SQL injection scenario.", False)])
    add_bullet(doc, [("xss_scenario.md — cross-site scripting scenario.", False)])


def build_team(doc):
    add_heading(doc, "19.  Team Structure & Roles", level=1)
    add_table(doc,
        ["Sub-Team", "Members", "Responsibility"],
        [
            ("Backend & AI", "Reem Amin (lead), Yousef Abdel Hady, Mohamed Shaban", "FastAPI, UnifiedRiskEngine, AI Advisory, Celery/Redis"),
            ("Frontend & Visualization", "Marize Ehap (lead), Omnia Helmy, Rahma Ebrahem", "React dashboard, D3.js topology, component architecture"),
            ("Security & Scanning", "Shahd Paher (lead), Mariz Ehap", "Nmap, Nuclei, OpenVAS, Wazuh, Elasticsearch"),
            ("DevOps & QA", "Omar Kapil (Team Leader + lead), Yosef Ali, Mazin Alla, Omar Tarek", "Docker, CI/CD, testing, documentation, presentation"),
        ])
    add_paragraph(doc, "Per-member deep-dives are maintained in team_roles/00_omar_kapil_role.md through 11_omar_tarek_role.md, each containing role summary, owned files, line-by-line code walkthrough, learning checklist and a week-by-week timeline in both English and Arabic.")
    add_paragraph(doc, "University presentation: July 2, 2026.")


def build_setup(doc):
    add_heading(doc, "20.  Setup, Installation & Deployment", level=1)

    add_heading(doc, "20.1  Prerequisites", level=2)
    add_bullet(doc, [("Docker Engine ≥ 24 and Docker Compose v2.", False)])
    add_bullet(doc, [("≥ 8 GB RAM, ≥ 20 GB free disk.", False)])
    add_bullet(doc, [("Optional: Google Gemini API key for full AI advisory.", False)])

    add_heading(doc, "20.2  Quick start", level=2)
    add_code_block(doc, """\
# 1) Clone
git clone <repo> && cd the-dashboard-project

# 2) Configure
cp .env.example .env
# Required keys: DATABASE_URL, REDIS_URL, JWT_SECRET, REPORT_SIGNING_KEY
# Optional:      GEMINI_API_KEY, AVATAR_UPLOAD_DIR, BACKEND_CORS_ORIGINS

# 3) Boot core stack
docker compose up -d --build

# 4) Boot lab (optional but recommended for first demo)
docker compose -f docker-compose.lab.yml up -d

# 5) Browse
#    Dashboard : http://localhost:5173
#    Swagger   : http://localhost:8000/docs
#    Kibana    : http://localhost:5601
#    n8n       : http://localhost:5678
""")

    add_heading(doc, "20.3  First login", level=2, red=True)
    add_paragraph(doc,
        "Default admin is seeded on first boot: admin@local / Admin#159. "
        "The system enforces force_password_change on first login — the "
        "operator MUST rotate the password before any other action.",
        red=True)

    add_heading(doc, "20.4  Migrations & seeding", level=2)
    add_paragraph(doc,
        "Alembic migrations run automatically at lifespan startup "
        "(alembic upgrade head). The orphan reaper sweeps any RUNNING "
        "scans left over from a previous boot. Runtime config overrides "
        "are loaded from the runtime_config table before the API begins "
        "accepting requests.", red=True)


def build_future(doc):
    add_heading(doc, "21.  Future Work", level=1)
    add_bullet(doc, [("Authenticated scanning — session cookie / OAuth crawl support.", False)])
    add_bullet(doc, [("Per-framework PDF templates (PCI-DSS, ISO 27001, NIST CSF) building on the control_tags already stored on every Finding.", True)])
    add_bullet(doc, [("Native SSO (OIDC / SAML) on top of the existing JWT layer.", True)])
    add_bullet(doc, [("Horizontal Celery autoscaling with per-target sharded queues.", False)])
    add_bullet(doc, [("Mobile companion app for KPI + critical-alert push.", False)])
    add_bullet(doc, [("Source-code SAST module bolted onto the existing AgentOrchestrator.", False)])
    add_bullet(doc, [("Air-gapped local LLM advisor (already abstracted behind ai_advisor.py).", True)])
    add_bullet(doc, [("Wazuh-agent auto-provisioning to close the last manual SIEM-integration step.", True)])


def build_conclusion(doc):
    add_heading(doc, "22.  Conclusion", level=1)
    add_paragraph(doc,
        "Orchestration Security Center demonstrates that an SME-grade "
        "security operations capability can be built entirely on "
        "open-source primitives, governed by deterministic rules, "
        "augmented (never directed) by AI, and operated by a "
        "non-specialist administrator from a single dashboard.")
    add_paragraph(doc,
        "This iteration moves the platform from a single-user prototype "
        "to a hardened multi-user system: a complete JWT + RBAC stack, "
        "tamper-evident audit chains, cryptographically signed reports, "
        "CVSS-aware risk breakdowns, finding deduplication, SLA tracking, "
        "framework control tagging, scope allowlists and per-target rate "
        "limits, scan checkpointing and orphan recovery, runtime "
        "feature-flag overrides, and a richer ten-container, "
        "four-subnet lab. With these foundations in place, the remaining "
        "Future Work items (authenticated scanning, native SSO, "
        "per-framework report templates, autoscaling, mobile) are "
        "incremental additions on top of a stable core.",
        red=True)


def build_changelog_appendix(doc):
    add_heading(doc, "Appendix A.  Changelog of Updates (Red Items)",
                level=1, red=True)
    add_paragraph(doc,
        "Every item below was marked in red in the body of this "
        "document because it exists in the live codebase but was NOT "
        "documented in the previous .md docs (PROJECT_OVERVIEW.md, "
        "USE_CASES_AND_EVALUATION.md, docs/API_GUIDE.md, "
        "docs/ARCHITECTURE_DIAGRAM.md).", red=True)

    # Group changelog by section
    grouped: dict[str, list[str]] = {}
    for section, item in CHANGELOG:
        grouped.setdefault(section, []).append(item)

    for section in grouped:
        add_heading(doc, section, level=2, red=True)
        for item in grouped[section]:
            add_paragraph(doc, "•  " + item, red=True, justify=False)


# ─────────────────────────────────────────────────────────────────────────────
# MAIN
# ─────────────────────────────────────────────────────────────────────────────

def main() -> Path:
    doc = Document()
    set_default_style(doc)
    add_page_numbers(doc)

    build_title_page(doc)
    build_abstract(doc)
    build_toc(doc)
    build_intro(doc)
    build_objectives(doc)
    build_background(doc)
    build_architecture(doc)
    build_tech_stack(doc)
    build_backend(doc)
    build_db_schema(doc)
    build_frontend(doc)
    build_auth_rbac(doc)
    build_lab(doc)
    build_pipeline(doc)
    build_scoring(doc)
    build_ai(doc)
    build_siem(doc)
    build_hardening(doc)
    build_use_cases(doc)
    build_pros_cons(doc)
    build_evidence(doc)
    build_team(doc)
    build_setup(doc)
    build_future(doc)
    build_conclusion(doc)
    build_changelog_appendix(doc)

    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    out = main()
    print(f"\n[OK] Wrote {out}")
    print(f"[OK] Total red-marked items: {len(CHANGELOG)}")
    print("\n=== Red changelog summary (grouped) ===")
    from collections import defaultdict
    grouped = defaultdict(list)
    for s, i in CHANGELOG:
        grouped[s].append(i)
    for section, items in grouped.items():
        print(f"\n[{section}]  ({len(items)} item{'s' if len(items) != 1 else ''})")
        for it in items:
            print(f"  - {it}")
